from docx import Document
from lxml import etree

from analisator import Parser

parser = Parser()

doc = Document()
p = doc.add_paragraph()
r = p.add_run()
formula_buf = '15'
r.text = r'<f>' + formula_buf + r'<\f>'
doc.save(f'test.docx')
