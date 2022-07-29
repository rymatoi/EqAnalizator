from docx import Document

doc = Document('in.docx')

doc_out = Document()
in_formula = False
p_out_just_started = False
formula_buf = ''
buf = ''
cur_del = ''
del_count = 0
del_type = None
p_out = None
for paragraph in doc.paragraphs:
    if buf and p_out:
        p_out.add_run(buf)
        buf = ''
    p_out = doc_out.add_paragraph()
    p_out_just_started = True
    for s in paragraph.text:
        if s == '$':
            cur_del += s
        else:
            if cur_del:
                del_count += 1
                del_type = cur_del
                cur_del = ''
            if del_count == 1:
                in_formula = True
                formula_buf += s
                if buf:
                    p_out.add_run(buf)
                    p_out_just_started = False
                    buf = ''
            elif del_count == 2:
                if not p_out_just_started:
                    p_out = doc_out.add_paragraph()
                in_formula = False
                del_count = 0
                if del_type == '$':
                    p_out.add_run(r'<f>' + formula_buf + r'</f>')
                elif del_type == '$$':
                    p_out.add_run(r'<f new-paragraph=true>' + formula_buf + r'</f>')
                else:
                    p_out.add_run(formula_buf)
                formula_buf = ''
                buf += s
            else:
                buf += s
    if in_formula and formula_buf:
        if not p_out_just_started:
            p_out = doc_out.add_paragraph()
        if del_type == '$':
            p_out.add_run(r'<f>' + formula_buf + r'</f>')
        elif del_type == '$$':
            p_out.add_run(r'<f new-paragraph=true new-paragraph-after=true>' + formula_buf + r'</f>')
        else:
            p_out.add_run(formula_buf)
        in_formula = False
        del_count = 0
        formula_buf = ''
        cur_del = ''
    else:
        if buf:
            p_out.add_run(buf)
            buf = ''

doc_out.save('out.docx')
