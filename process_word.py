import docx
from bs4 import BeautifulSoup
from lxml import etree
import json

from analisator import Parser


def process_word(word_in, word_out):
    parser = Parser()
    doc = docx.Document(word_in)
    rdoc = docx.Document()
    rp = None
    new_paragraph_after = False

    for paragraph in doc.paragraphs:
        print(paragraph.text)
        soup = BeautifulSoup(paragraph.text, 'html.parser')
        formulas = soup.findAll('f')
        if formulas:
            for formula in formulas:
                rendered_formula = parser.render_equation_(formula.text)
                formula_xml = etree.fromstring(rendered_formula)
                new_paragraph = json.loads(formula.attrs.get("new-paragraph", 'false').lower())
                new_paragraph_after = json.loads(formula.attrs.get("new-paragraph-after", 'false').lower())
                if new_paragraph:
                    rp = rdoc.add_paragraph()
                    rp._element.append(formula_xml)
                else:
                    if not rp:
                        rp = rdoc.add_paragraph()
                    rp._element.append(formula_xml)

        else:
            if not rp:
                rp = rdoc.add_paragraph()
            if new_paragraph_after:
                rp = rdoc.add_paragraph()
            for el in paragraph._element:
                rp._element.append(el)
            rp = None

    rdoc.save(word_out)
